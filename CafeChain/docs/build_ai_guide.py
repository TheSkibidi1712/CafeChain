from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "CafeChain_AI_Business_And_User_Guide.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = value
        shade(cell, "E8EEF5")
        set_cell_margins(cell)
        if widths:
            set_col_width(cell, widths[i])
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if widths:
                set_col_width(cell, widths[i])
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_bullets(doc, values, numbered=False):
    for value in values:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(value)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = "Code"
    p.add_run(text)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(5)
styles["Normal"].paragraph_format.line_spacing = 1.12
for name, size, color in [("Title", 28, "17365D"), ("Heading 1", 18, "17365D"), ("Heading 2", 13, "2F5597"), ("Heading 3", 11, "404040")]:
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(color)
    styles[name].font.bold = True
    styles[name].paragraph_format.space_before = Pt(10 if name != "Title" else 0)
    styles[name].paragraph_format.space_after = Pt(5)
if "Code" not in styles:
    code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = styles["Code"]
code_style.font.name = "Consolas"
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor.from_string("404040")
code_style.paragraph_format.left_indent = Inches(0.25)
code_style.paragraph_format.space_after = Pt(4)

# Editorial cover header/footer
header = section.header.paragraphs[0]
header.text = "CAFÉCHAIN  /  AI BUSINESS & USER GUIDE"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs:
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("6B7280")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Internal reference • ")
add_page_field(footer)
for r in footer.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("6B7280")

cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(44)
cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = cover.add_run("CAFÉCHAIN")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string("2F5597")
title = doc.add_paragraph(style="Title")
title.add_run("AI Business\n& User Guide")
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_before = Pt(8)
subtitle.add_run("MVC AI Dashboard • Inventory/Reorder AI • AI Image • SignalR notifications").italic = True
meta = doc.add_paragraph()
meta.paragraph_format.space_before = Pt(40)
meta.add_run("Purpose\n").bold = True
meta.add_run("A maintenance and user reference for the current CaféChain AI architecture.\n")
meta.add_run("Scope\n").bold = True
meta.add_run("Read-only analytics, evidence-grounded explanations, inventory alerts and image assistance.\n")
meta.add_run("Security note\n").bold = True
meta.add_run("Secrets are supplied through environment variables/User Secrets; no API key is stored in this guide.")
doc.add_page_break()

add_heading(doc, "1. Tổng quan hệ thống AI", 1)
doc.add_paragraph("CaféChain hiện dùng một kiến trúc AI có kiểm soát: dữ liệu được lấy qua repository/service, rule engine phát hiện tín hiệu chắc chắn, rồi Ollama (khi sẵn sàng) tạo phần giải thích có cấu trúc. Dashboard AI chạy trong ASP.NET Core MVC Admin; React chỉ đảm nhiệm chuông notification realtime.")
add_table(doc, ["Thành phần", "Mục đích", "Nguồn/đường chạy"], [
    ("AI Dashboard", "Phân tích doanh thu, đơn, sản phẩm, cửa hàng, kho và NCC; trả fact/inference/statistic.", "MVC Dashboard → DashboardIntelligenceService → DashboardRepository → Ollama/fallback"),
    ("Inventory/Reorder AI", "Tính consumption, threshold, shortage risk và gợi ý đặt hàng.", "Rule/statistic engine + repository; notification dùng StaffNotification"),
    ("Supplier/Price AI", "So sánh chất lượng NCC, giá nhập, spend, issue và PO quá hạn.", "SupplierIntelligenceService + dataset read-only"),
    ("AI Image", "Tạo query, tìm ảnh Pexels, chấm điểm; fallback ComfyUI.", "AIImagePipelineService → Pexels → ComfyUI"),
    ("Forecast/Anomaly/POS recommendation", "Các AI chuyên biệt hiện có, dùng rule và thống kê khi đủ dữ liệu.", "Application AI services; không thay đổi transaction"),
])

add_heading(doc, "2. Kiến trúc AI", 1)
add_code(doc, "Frontend MVC → Controller → AI/Application Service → Data Service/Repository → Prompt + Skill → Ollama provider → Structured result → MVC")
add_bullets(doc, [
    "Controller chỉ nhận/validate request, kiểm tra antiforgery và gọi service.",
    "Service lập data plan cố định theo intent; AI không được chọn SQL, table, column hay stored procedure.",
    "Repository chỉ đọc dữ liệu qua query/stored procedure đã kiểm soát; Dashboard mặc định read-only.",
    "Skill và schema quy định grounding, evidence, fact/inference và xử lý dữ liệu thiếu.",
    "Nếu Ollama offline, timeout hoặc JSON sai, service trả deterministic structured fallback."
])

add_heading(doc, "3. AI Dashboard MVC", 1)
doc.add_paragraph("Người dùng đặt câu hỏi tự do trong Dashboard Admin MVC. Intent parser ưu tiên Ollama; parser từ khóa chỉ là fallback. Server ánh xạ intent sang dataset allowlist, kiểm tra StoreScope rồi mới query.")
add_table(doc, ["Intent", "Dataset chính"], [
    ("REVENUE_ANALYSIS / SALES_TREND", "Net sales trend, store ranking, category/product performance, heatmap"),
    ("ORDER_ANALYSIS", "Order status/cancellation, hourly orders, payment mix"),
    ("PRODUCT_PERFORMANCE", "Product/category, volume-margin, size và topping"),
    ("STORE_COMPARISON", "Store ranking, cancellation và inventory risk"),
    ("INVENTORY_ANALYSIS / REORDER_ANALYSIS", "Shortage/threshold, consumption, waste, reorder"),
    ("SUPPLIER_ANALYSIS", "Supplier quality, prices, spend, issues, overdue POs"),
    ("ANOMALY_DETECTION", "Operational alerts và baseline/current datasets"),
    ("GENERAL_BUSINESS_SUMMARY / STATISTICS_REQUEST", "Dataset tương ứng focus metric được allowlist"),
])
add_heading(doc, "Structured output và evidence", 2)
add_code(doc, '{ "analysisId": "...", "intent": "REVENUE_ANALYSIS", "dataPeriod": {}, "dataStatus": "Complete|Partial|Insufficient", "summary": "...", "facts": [], "inferences": [], "statistics": [], "anomalies": [], "recommendations": [], "confidence": 0.0, "charts": [], "warnings": [], "aiStatus": "Available|Fallback|Offline" }')
add_bullets(doc, [
    "Fact, statistic, anomaly và chart data do server tạo từ dataset; mỗi evidence có evidenceId, nguồn, kỳ và giá trị gốc.",
    "AI chỉ tạo summary, inference và recommendation; narrative phải tham chiếu evidence đã cung cấp.",
    "Validator loại bỏ field ngoài schema và nhận định không có evidence. Không đủ dữ liệu phải ghi rõ: “Không đủ dữ liệu để kết luận.”",
    "Rule engine vẫn giữ: doanh thu giảm 20%, hủy tăng 30%, top seller giảm 40%, stock dưới min threshold, waste/supplier risk. Rule là tín hiệu, không phải template trả lời."
])
add_heading(doc, "Câu hỏi tự do mẫu", 2)
add_bullets(doc, ["Tại sao doanh thu hôm nay giảm?", "Chi nhánh nào đang hoạt động kém?", "So sánh doanh thu tuần này và tuần trước.", "Có nguyên liệu nào sắp thiếu không?", "Tôi nên chú ý điều gì hôm nay?"])

add_heading(doc, "4. Inventory/Reorder AI", 1)
doc.add_paragraph("Inventory AI dùng tồn hiện tại, min threshold, consumption theo kỳ, waste và lịch sử xuất bán để phát hiện shortage risk và tạo reorder suggestion. Khi POS trừ kho, InventoryDeductionService kích hoạt StockAlertService; POSOrderService không bị sửa.")
add_bullets(doc, [
    "Notification dùng StaffNotification hiện có, không tạo subsystem thứ hai.",
    "Audience cần Permission Notification.View và StoreScope; group SignalR có dạng store:{storeId}:permission:Notification.View.",
    "Deduplication key gồm RecipientStaffId + StoreId + NotificationType + EntityType + EntityId.",
    "Cooldown mặc định 15 phút: tín hiệu trùng trong cooldown không tạo toast mới; escalation hoặc sau cooldown mới push lại. Resolve đóng incident; tái diễn tạo incident mới."
])

add_heading(doc, "5. AI Supplier/Price", 1)
doc.add_paragraph("Supplier intelligence đọc chất lượng NCC, giá theo lịch sử, package quantity/unit, tổng spend, issue mix và PO quá hạn. Recommendation chỉ được nêu khi có evidence về giá, chất lượng hoặc thời gian giao; không tự suy diễn chi phí nếu nguồn cost không hợp lệ.")

add_heading(doc, "6. AI Image, Pexels và ComfyUI", 1)
add_code(doc, "AI Suggestion → Search query → Pexels → metadata/match validation → fallback ComfyUI → generated image")
add_bullets(doc, [
    "Pexels được ưu tiên khi có API key và ảnh đạt kích thước/match score tối thiểu.",
    "Nếu không có ảnh phù hợp hoặc Pexels lỗi, ComfyUI dùng workflow txt2img/img2img trong Resources/AI/ComfyUI.",
    "Positive prompt được tạo từ visual specification; negative prompt mặc định loại text, logo, watermark, người, blur và vật thể méo.",
    "Checkpoint/workflow, kích thước, sampler, steps và CFG đọc từ cấu hình; không ghi key thật trong source hoặc tài liệu."
])

add_heading(doc, "7. Ollama", 1)
add_bullets(doc, [
    "Cài Ollama từ trang chính thức, khởi động service và kiểm tra endpoint mặc định http://localhost:11434.",
    "Model hiện cấu hình là qwen3:4b; kiểm tra bằng `ollama list`, tải bằng `ollama pull qwen3:4b`.",
    "CaféChain dùng timeout 120 giây, keep-alive 5m, temperature thấp để tăng tính ổn định JSON.",
    "Test nhanh: `ollama run qwen3:4b` rồi gửi một câu hỏi ngắn; Dashboard sẽ hiển thị aiStatus Available khi gọi thành công.",
    "Khi Ollama offline/timeout/JSON invalid, fallback server vẫn trả facts/statistics/rules và nêu rõ aiStatus."
])

add_heading(doc, "8. ComfyUI", 1)
add_bullets(doc, [
    "Cài ComfyUI, đặt checkpoint đúng tên trong cấu hình, chạy port 8188 và mở giao diện để test workflow.",
    "Workflow hiện dùng product-txt2img.json và product-img2img.json; checkpoint node 4, sampler node 3, output node 9.",
    "Cấu hình mặc định: 512×512, 28 steps, Euler/normal, CFG 7; timeout 180 giây.",
    "Nếu offline: kiểm tra process/port 8188, checkpoint, node id và log queue. Pipeline giữ fallback Pexels nếu online source còn ảnh hợp lệ."
])

add_heading(doc, "9. Pexels", 1)
doc.add_paragraph("API key không nằm trong tracked configuration. Cấp qua `Pexels__ApiKey` hoặc User Secrets, ví dụ `dotnet user-secrets set \"Pexels:ApiKey\" \"<your-key>\"`. Không ghi key vào commit, log hay tài liệu.")
add_bullets(doc, ["Query do AI image suggestion tạo; server giới hạn số query/page.", "Metadata scorer kiểm tra tên, mô tả, kích thước và match score.", "Không đạt ngưỡng hoặc API lỗi thì fallback ComfyUI; thiếu key thì bỏ qua Pexels an toàn."])

add_heading(doc, "10. SignalR Notification", 1)
add_code(doc, "POS → InventoryDeduction/StockAlert → StaffNotification upsert → InventoryNotificationHub (/hubs/inventory-notifications) → CafeChain.Frontend notification bell")
add_bullets(doc, [
    "Hub chỉ xác thực JWT, kiểm tra store claim/permission và quản lý group; business logic nằm ở delivery service.",
    "Publisher chỉ chạy sau SaveChanges thành công. Event có eventId, storeId, type, severity, changeKind, entity và shouldToast.",
    "Frontend dùng accessTokenFactory, automatic reconnect, dedupe eventId, cập nhật REST list/badge và giữ polling fallback khi offline.",
    "Nhân viên khác store hoặc thiếu Notification.View không join group và không nhận event."
])

add_heading(doc, "11. Hướng dẫn sử dụng AI Dashboard", 1)
add_bullets(doc, [
    "Bước 1: Mở Admin Dashboard trong MVC.",
    "Bước 2: Chọn cửa hàng/phạm vi StoreScope và khoảng thời gian.",
    "Bước 3: Nhập câu hỏi tự do, ví dụ “So sánh doanh thu tuần này và tuần trước”.",
    "Bước 4: Bấm Analyze và chờ trạng thái loading.",
    "Bước 5: Đọc Summary, Fact/Statistic, Inference, Anomaly và Recommendation tách biệt.",
    "Bước 6: Kiểm tra period, evidence và cảnh báo Partial/Insufficient trước khi ra quyết định."
])

add_heading(doc, "12. Troubleshooting", 1)
add_table(doc, ["Lỗi", "Nguyên nhân thường gặp", "Kiểm tra", "Xử lý"], [
    ("Ollama không chạy", "Service/port 11434 offline", "Mở endpoint hoặc `ollama list`", "Khởi động Ollama; dùng fallback nếu chưa sẵn sàng"),
    ("Không tìm thấy model", "Chưa pull qwen3:4b", "`ollama list`", "Pull đúng model và khớp cấu hình"),
    ("Timeout", "Model nặng hoặc máy bận", "Log request/timeout", "Tăng tài nguyên hoặc giữ fallback; không retry vô hạn"),
    ("JSON AI invalid", "Model thêm markdown/field lạ", "AI response log đã sanitize", "Schema validator loại bỏ; trả structured fallback"),
    ("Pexels không có ảnh phù hợp", "Query/match score thấp hoặc thiếu key", "Kiểm tra Pexels config", "Điều chỉnh query hoặc dùng ComfyUI"),
    ("ComfyUI offline", "Port/checkpoint/workflow sai", "Kiểm tra port 8188 và queue", "Khởi động lại, kiểm tra node/checkpoint"),
    ("SignalR disconnect", "JWT hết hạn/mạng chập chờn", "Console connection state", "Automatic reconnect; polling là fallback"),
    ("Notification không realtime", "Hub route/JWT/group lỗi", "Kiểm tra /hubs/inventory-notifications", "Kiểm tra token, store claim, publisher sau SaveChanges"),
    ("Không nhận do permission", "Thiếu Notification.View hoặc ngoài StoreScope", "Kiểm tra policy/scope", "Cấp đúng permission/store; không hard-code role"),
    ("AI Dashboard không có dữ liệu", "Kỳ không có order/warehouse data", "Kiểm tra period và dataStatus", "Chọn kỳ có dữ liệu; không suy đoán khi Insufficient"),
])

doc.add_paragraph("Bảo trì: mọi thay đổi Dashboard phải giữ allowlist dataset, structured schema, StoreScope và read-only boundary. Không thêm AI Dashboard route/component vào CafeChain.Frontend; không chỉnh các form Kho & Cung ứng ngoài chuông notification.")
doc.save(OUT)
print(str(OUT).encode("ascii", "backslashreplace").decode("ascii"))
